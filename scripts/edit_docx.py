#!/usr/bin/env python3
"""
Surgical find/replace inside a .docx, preserving comments, tracked changes,
and all surrounding formatting.

Use this instead of re-inserting a section once a reviewer has started
commenting or editing — md_to_docx.py replaces the whole section and would
discard their review artifacts.

Usage:
    edit_docx.py <file.docx> --replace "old text" "new text" [--replace ...] [--dry-run]
    edit_docx.py <file.docx> --from-file edits.tsv        # one "old<TAB>new" per line

Matching is on visible paragraph text (including inside tables). Replacement is
applied to the first run containing the match and the remaining runs of the match
are cleared, so the new text inherits the formatting of the run it replaces.
"""
import argparse
import sys

import docx


def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nt in cell.tables:
                    for r2 in nt.rows:
                        for c2 in r2.cells:
                            for p in c2.paragraphs:
                                yield p


def replace_in_paragraph(p, old, new):
    """Replace `old` with `new` across run boundaries. Returns True if changed."""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    idx = full.index(old)
    end = idx + len(old)

    pos = 0
    start_run = None
    for r in p.runs:
        rlen = len(r.text)
        rstart, rend = pos, pos + rlen
        if rend <= idx or rstart >= end:
            pos = rend
            continue
        # portion of this run covered by the match
        local_start = max(0, idx - rstart)
        local_end = min(rlen, end - rstart)
        if start_run is None:
            start_run = r
            r.text = r.text[:local_start] + new + r.text[local_end:]
        else:
            r.text = r.text[:local_start] + r.text[local_end:]
        pos = rend
    return start_run is not None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--replace", nargs=2, action="append", metavar=("OLD", "NEW"), default=[])
    ap.add_argument("--from-file", help="TSV file: old<TAB>new per line")
    ap.add_argument("--dry-run", action="store_true")
    a = ap.parse_args()

    pairs = list(a.replace)
    if a.from_file:
        for line in open(a.from_file, encoding="utf-8"):
            line = line.rstrip("\n")
            if not line.strip() or line.lstrip().startswith("#"):
                continue
            old, _, new = line.partition("\t")
            pairs.append((old, new))
    if not pairs:
        sys.exit("nothing to do — pass --replace or --from-file")

    doc = docx.Document(a.docx)
    counts = {}
    for old, new in pairs:
        n = 0
        for p in iter_paragraphs(doc):
            while replace_in_paragraph(p, old, new):
                n += 1
                if not a.dry_run:
                    continue
                break
        counts[old] = n

    for old, n in counts.items():
        mark = "ok " if n else "!! "
        print("%s%d match(es)  %r" % (mark, n, old[:72]))

    missing = [o for o, n in counts.items() if n == 0]
    if a.dry_run:
        print("\n(dry run — nothing written)")
    else:
        doc.save(a.docx)
        print("\nsaved %s" % a.docx)
    if missing:
        print("WARNING: %d pattern(s) not found — check wording/typography (en dash, curly quotes)"
              % len(missing))


if __name__ == "__main__":
    main()
