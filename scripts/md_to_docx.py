#!/usr/bin/env python3
"""
Insert a drafted DSE section (markdown) into the official Word package template.

Usage:
    md_to_docx.py <source.md> <target.docx> --after "<Heading 1 text>" --before "<Heading 1 text>"

Content is inserted between the two anchor Heading 1 paragraphs, replacing anything
already there. Repo-internal scaffolding (scope notes, vault paths, harvest logs,
open items, "Explicitly Excluded" tables) is filtered out — see SKIP_SECTIONS.

`[verify]` / `[pending ...]` markers are rendered in red so they are obvious during
final cleanup.
"""
import re
import sys
import argparse

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Heading + everything under it is dropped ----
SKIP_SUBTREE = (
    "explicitly excluded",
    "vault harvest log",
    "open items",
    "gaps / bruce",
    "note on page budget",
    "original xls",
    "updated xls",
)

# ---- Only the heading line is dropped; its content is promoted ----
SKIP_HEADING_ONLY = (
    "draft package body",
    "flagship innovations",
)

# ---- Standalone lines dropped ----
# Cross-references are kept — they are the gold-standard package pattern.
# Evidence/Vault lines are dropped: they carry Obsidian paths meaningless to a reviewer.
SKIP_LINE_PAT = re.compile(
    r"^\s*(\*\*Suggested package length|\*\*Vault:|\*\*Evidence:|\*\*Source:|"
    r"\*Account-level placeholders|\*Dollar figures marked|\*Letters stored|"
    r"\*Session preparation|\*Exec thru-line|\*From vault|\*Original SMART)",
    re.I,
)

VERIFY_PAT = re.compile(r"(\[(?:verify|pending|confirm|add |Bruce decision|awaiting|extend|scores|full session|attendee|source)[^\]]*\])", re.I)


# ------------------------------------------------------------------ inline runs
def add_hyperlink(paragraph, url, text, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag in ("w:b",) if bold else ():
        rPr.append(OxmlElement(tag))
    if italic:
        rPr.append(OxmlElement("w:i"))
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


TOKEN = re.compile(
    r"(\[[^\]]+\]\([^)]+\)"      # link
    r"|\*\*[^*]+\*\*"            # bold
    r"|`[^`]+`"                  # code
    r"|\*[^*\n]+\*)"             # italic
)


def emit_runs(p, text):
    """Render inline markdown into runs on paragraph p."""
    for chunk in filter(None, TOKEN.split(text)):
        m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
        if m:
            label, url = m.group(1), m.group(2)
            label = re.sub(r"\*\*|`", "", label)
            if url.startswith(("http://", "https://")):
                add_hyperlink(p, url, label)
            else:
                p.add_run(label)
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            p.add_run(chunk[2:-2]).bold = True
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            # Backticks in the source mark repo and org names, not code. Render them
            # in the body font: the official rules require CiscoSansTT throughout,
            # and a monospace switch mid-sentence reads as a formatting error.
            p.add_run(chunk[1:-1])
            continue
        if chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            p.add_run(chunk[1:-1]).italic = True
            continue
        # plain text — split out [verify] markers so they can be colored
        for piece in filter(None, VERIFY_PAT.split(chunk)):
            r = p.add_run(piece)
            if VERIFY_PAT.fullmatch(piece):
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                r.bold = True


# ------------------------------------------------------------------ block parse
def parse_blocks(md, title=None):
    """Yield ('h2'|'h3'|'h4'|'p'|'ul'|'table', payload).

    `title` is the section name the template already provides; an H2 matching it
    exactly is dropped so the heading is not duplicated.
    """
    lines = md.split("\n")
    i, blocks, skipping = 0, [], False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2).strip()
            low = text.lower()
            if any(s in low for s in SKIP_HEADING_ONLY):
                i += 1
                continue
            if any(s in low for s in SKIP_SUBTREE):
                skip_level = level
                i += 1
                while i < len(lines):
                    nm = re.match(r"^(#{1,4})\s+(.*)$", lines[i].strip())
                    if nm and len(nm.group(1)) <= skip_level:
                        break
                    i += 1
                continue
            if level == 1:
                # A level-1 heading is the document title; the template already
                # carries it as Heading 1. Drop it, or demote if it differs.
                if not title or text.strip().lower() == title.strip().lower():
                    i += 1
                    continue
                level = 2
            if level == 2 and title and text.strip().lower() == title.strip().lower():
                i += 1
                continue  # section title — the template already carries it
            blocks.append(("h%d" % level, text))
            i += 1
            continue

        # Reviewer comments: line-initial `//` or an HTML comment.
        # Line-initial only, so `https://…` inside prose is never eaten.
        if stripped.startswith("//") or (
            stripped.startswith("<!--") and stripped.endswith("-->")
        ):
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Blockquote. These carry the leader-assessment and customer quotes, so they
        # must render -- an earlier version dropped them along with the horizontal
        # rules and silently lost five Brook Crossman quotes from the package.
        if stripped.startswith(">"):
            # One quote per line. Consecutive '>' lines in these files are separate
            # attributed quotes (e.g. four manager assessments in a row), not one
            # wrapped paragraph, so joining them would merge distinct quotes.
            while i < len(lines) and lines[i].strip().startswith(">"):
                text = re.sub(r"^\s*>\s?", "", lines[i].rstrip()).strip()
                if text:
                    blocks.append(("quote", text))
                i += 1
            continue
        if not stripped:
            i += 1
            continue
        if SKIP_LINE_PAT.match(stripped):
            i += 1
            continue

        # table
        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        # bullets
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ul", items))
            continue

        # paragraph (join wrapped lines)
        para = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            # A line opening with bold is always a new labeled block
            # (e.g. the four "Financial / Competitive / Strategic / Customer impact" lines),
            # so do not fold it into the previous paragraph.
            if (not nxt or nxt.startswith(("|", ">", "#", "---", "**"))
                    or re.match(r"^[-*]\s+", nxt)):
                break
            para.append(nxt)
            i += 1
        blocks.append(("p", " ".join(para)))
    return blocks


# ------------------------------------------------------------------ write
def build(doc, blocks, anchor):
    """Create blocks and move them immediately before `anchor` paragraph element."""
    def place(el):
        anchor.addprevious(el)

    for kind, payload in blocks:
        if kind in ("h2", "h3", "h4"):
            style = {"h2": "Heading 2", "h3": "Heading 3", "h4": "Heading 4"}[kind]
            p = doc.add_paragraph(style=style)
            emit_runs(p, payload)
            place(p._p)

        elif kind == "p":
            p = doc.add_paragraph()
            emit_runs(p, payload)
            place(p._p)

        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            emit_runs(p, payload)
            place(p._p)

        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Paragraph")
                p.paragraph_format.left_indent = Inches(0.35)
                p.paragraph_format.first_line_indent = Inches(-0.18)
                p.paragraph_format.space_after = Pt(2)
                p.add_run("•  ")
                emit_runs(p, item)
                place(p._p)

        elif kind == "table":
            ncols = max(len(r) for r in payload)
            t = doc.add_table(rows=0, cols=ncols)
            t.style = "Table Grid"
            t.autofit = True
            for ri, row in enumerate(payload):
                cells = t.add_row().cells
                for ci in range(ncols):
                    text = row[ci] if ci < len(row) else ""
                    cell = cells[ci]
                    cp = cell.paragraphs[0]
                    cp.paragraph_format.space_before = Pt(1)
                    cp.paragraph_format.space_after = Pt(1)
                    emit_runs(cp, text)
                    for r in cp.runs:
                        r.font.size = Pt(10)   # official minimum: no smaller than 10 pt
                        if ri == 0:
                            r.bold = True
            place(t._tbl)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
            place(spacer._p)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("source")
    ap.add_argument("target")
    ap.add_argument("--after", required=True, help="Heading 1 text to insert after")
    ap.add_argument("--before", required=True, help="Heading 1 text to insert before")
    ap.add_argument("--title", help="Markdown H2 title to drop (defaults to --after)")
    a = ap.parse_args()

    doc = docx.Document(a.target)
    paras = doc.paragraphs
    to_end = a.before.strip().upper() == "END"

    def find(text):
        # Match Heading 1 only. Inserted section content contains Heading 2s such as
        # "Global Impact Summary" that would otherwise capture the anchor.
        for idx, p in enumerate(paras):
            if p.style.name in ("Heading 1", "Heading1") and \
               p.text.strip().rstrip().lower().startswith(text.lower()):
                return idx
        raise SystemExit("anchor not found as a Heading 1: %r" % text)

    start_el = paras[find(a.after)]._p
    if to_end:
        # append before the trailing sectPr so content lands at end of the body
        end_el = doc.element.body.find(qn("w:sectPr"))
        if end_el is None:
            end_el = doc.add_paragraph()._p
    else:
        end_el = paras[find(a.before)]._p

    # Clear everything already between the anchors — paragraphs AND tables.
    # Iterating doc.paragraphs alone leaves stale tables behind, which then
    # accumulate on every re-run.
    body = doc.element.body
    children = list(body.iterchildren())
    try:
        i0, i1 = children.index(start_el), children.index(end_el)
    except ValueError:
        raise SystemExit("could not locate anchors in document body")
    if i1 <= i0:
        raise SystemExit("--before must come after --after in the document")

    removed = {"p": 0, "tbl": 0}
    for el in children[i0 + 1:i1]:
        if el.tag.endswith("}p"):
            text = "".join(el.itertext()).strip().lower()
            if text.startswith(("suggested length", "no page limit")):
                continue
            removed["p"] += 1
        elif el.tag.endswith("}tbl"):
            removed["tbl"] += 1
        else:
            continue
        body.remove(el)

    anchor = end_el

    md = open(a.source, encoding="utf-8").read()
    blocks = parse_blocks(md, title=a.title or a.after)
    build(doc, blocks, anchor)
    doc.save(a.target)

    kinds = {}
    for k, _ in blocks:
        kinds[k] = kinds.get(k, 0) + 1
    print("inserted into %r" % a.target)
    print("  cleared %d paragraph(s), %d table(s)" % (removed["p"], removed["tbl"]))
    print("  blocks: " + ", ".join("%s=%d" % kv for kv in sorted(kinds.items())))


if __name__ == "__main__":
    main()
