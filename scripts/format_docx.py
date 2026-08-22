#!/usr/bin/env python3
"""
Apply the DSE package's page-layout and typography rules to a .docx.

The official guidelines constrain only: <=50 pages (excluding index and appendix),
font no smaller than 10 pt, single spaced, CiscoSansTT, working hyperlinks.
**Margins are not specified**, so they are ours to set.

What this does:
  * page margins (default 0.75in all round; template ships with 1.25in left/right)
  * Normal to 10pt CiscoSansTT, single spaced, tight space-after
  * heading fonts and reduced space-before
  * table column widths proportional to content, replacing equal-width columns
  * tighter table cell margins
  * removes the template's per-section guidance lines
    ("Suggested length: 3-4 pages", "No page limit")

Usage:
    format_docx.py <file.docx> [--margin 0.75] [--body-pt 10] [--keep-guidance]
"""
import argparse
import re

import docx
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GUIDANCE_PREFIXES = ("suggested length", "no page limit")
FONT = "CiscoSansTT"


def set_margins(doc, inches):
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(inches)
        s.top_margin = s.bottom_margin = Inches(inches)


def style_fonts(doc, body_pt):
    n = doc.styles["Normal"]
    n.font.name = FONT
    n.font.size = Pt(body_pt)
    pf = n.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    # East-Asian / complex-script font names, so Word does not substitute
    rpr = n.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)

    for name, size, before in (("Heading 1", body_pt + 5, 10),
                               ("Heading 2", body_pt + 3, 8),
                               ("Heading 3", body_pt + 1, 6),
                               ("Heading 4", body_pt, 5)):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.size = Pt(size)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.line_spacing = 1.0

    for name in ("List Paragraph", "Body Text"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = FONT
        st.font.size = Pt(body_pt)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.space_after = Pt(2)


def tighten_cell_margins(table, twips=54):
    """Default is 108 twips (0.075in) left/right. Halve it."""
    tblPr = table._tbl.tblPr
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar"); tblPr.append(mar)
    for side, val in (("left", twips), ("right", twips), ("top", 0), ("bottom", 0)):
        el = mar.find(qn("w:" + side))
        if el is None:
            el = OxmlElement("w:" + side); mar.append(el)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")


def size_columns(table, usable_in, min_in=0.55):
    """Allocate column widths proportional to content, damped so one long cell
    cannot swallow the table. Replaces python-docx's equal-width default."""
    rows = table.rows
    if not rows:
        return
    ncol = len(table.columns)
    scores = []
    for j in range(ncol):
        longest = 0
        for r in rows:
            try:
                longest = max(longest, len(r.cells[j].text.strip()))
            except IndexError:
                pass
        scores.append(max(longest, 1))
    # damp: ^0.62 keeps a 400-char cell from being 40x a 10-char cell
    w = [s ** 0.62 for s in scores]
    total = sum(w)
    widths = [max(min_in, usable_in * x / total) for x in w]
    # renormalize after applying the floor
    scale = usable_in / sum(widths)
    widths = [x * scale for x in widths]

    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for j, win in enumerate(widths):
        for r in rows:
            try:
                r.cells[j].width = Inches(round(win, 3))
            except IndexError:
                pass


def set_toc_levels(doc, lo=1, hi=2):
    """Limit the table of contents to heading levels lo..hi.

    Two things have to change. The field switch `TOC \\o "1-4"` governs what Word
    builds on refresh, but the document also carries a *cached* rendering of the
    old TOC (one paragraph per entry, styled TOC1..TOC4). Word shows the cache
    until the field is refreshed, so the dropped levels have to be removed too or
    they keep occupying pages.
    """
    changed_field = False
    for el in doc.element.body.iter(qn("w:instrText")):
        if el.text and "TOC" in el.text and "\\o" in el.text:
            new = re.sub(r'\\o\s*"\d+-\d+"', '\\\\o "%d-%d"' % (lo, hi), el.text)
            if new != el.text:
                el.text = new
                changed_field = True

    # Walk every w:p in the tree, not doc.paragraphs: Word wraps the TOC in a
    # content control (w:sdt), and doc.paragraphs only sees direct body children.
    dropped = 0
    for p in list(doc.element.body.iter(qn("w:p"))):
        pStyle = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pStyle is None:
            continue
        val = (pStyle.get(qn("w:val")) or "").replace(" ", "")
        m = re.fullmatch(r"TOC(\d+)", val, re.I)
        if m and not (lo <= int(m.group(1)) <= hi):
            p.getparent().remove(p)
            dropped += 1
    return changed_field, dropped


def normalize_run_fonts(doc, body_pt):
    """Force every run to the required font.

    Earlier builds rendered markdown backtick spans in Consolas. The official
    rules require CiscoSansTT throughout, and a monospace switch mid-sentence
    reads as a formatting error. Run-level rFonts overrides beat the Normal
    style, so they have to be cleared here rather than at the style level.
    """
    fixed = 0
    for r in doc.element.body.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        rf = rPr.find(qn("w:rFonts"))
        if rf is None:
            continue
        if (rf.get(qn("w:ascii")) or "") != FONT:
            for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rf.set(qn(a), FONT)
            fixed += 1
    return fixed


def strip_guidance(doc):
    removed = 0
    for p in list(doc.paragraphs):
        t = p.text.strip().lower()
        if t.startswith(GUIDANCE_PREFIXES):
            p._p.getparent().remove(p._p)
            removed += 1
    return removed


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--margin", type=float, default=0.75)
    ap.add_argument("--body-pt", type=float, default=10)
    ap.add_argument("--keep-guidance", action="store_true")
    ap.add_argument("--toc-levels", default="1-2",
                    help='heading levels to keep in the TOC, e.g. "1-2"')
    a = ap.parse_args()

    doc = docx.Document(a.docx)
    before = doc.sections[0]
    old = "%.2f x %.2f in" % (
        before.page_width.inches - before.left_margin.inches - before.right_margin.inches,
        before.page_height.inches - before.top_margin.inches - before.bottom_margin.inches)

    set_margins(doc, a.margin)
    style_fonts(doc, a.body_pt)

    s = doc.sections[0]
    usable = s.page_width.inches - s.left_margin.inches - s.right_margin.inches
    for t in doc.tables:
        tighten_cell_margins(t)
        size_columns(t, usable)

    font_fixed = normalize_run_fonts(doc, a.body_pt)

    lo, hi = (int(v) for v in a.toc_levels.split("-"))
    toc_field, toc_dropped = set_toc_levels(doc, lo, hi)

    removed = 0 if a.keep_guidance else strip_guidance(doc)
    doc.save(a.docx)

    new = "%.2f x %.2f in" % (usable, s.page_height.inches - s.top_margin.inches - s.bottom_margin.inches)
    gain = ((usable) / (before.page_width.inches - 2 * 1.25)) if False else None
    print("text area : %s  ->  %s" % (old, new))
    print("body font : %.0f pt %s, single spaced, 3pt after" % (a.body_pt, FONT))
    print("tables    : %d re-sized to content-proportional columns; cell margins halved" % len(doc.tables))
    print("fonts     : %d run(s) forced to %s (was Consolas or other)" % (font_fixed, FONT))
    print("guidance  : %d 'Suggested length' / 'No page limit' lines removed" % removed)
    print("toc       : levels %d-%d%s; %d cached out-of-range entries dropped"
          % (lo, hi, " (field switch updated)" if toc_field else "", toc_dropped))
    print("saved     : %s" % a.docx)


if __name__ == "__main__":
    main()
